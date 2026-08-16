"""Built-in document tools: read_pdf, read_docx — read-only extraction.

Same discipline as the plain file tools: one FilesystemEffect declared
from params, execution only on the gate-blessed path via the shared
no-follow open, zero permission logic. Extracted text is returned as a
tool result — untrusted DATA under Durable Rule 12, exactly like
read_file's output.

Bounded extraction: a hard input-size guard refuses absurd files before
parsing, and text accumulates page-by-page (PDF) / block-by-block (docx)
stopping AT the cap — a large document is truncated with a clear marker
during extraction, never loaded into the conversation whole.

The parsing libraries (pypdf, python-docx) are sealed inside this module.
No OCR: an image-only PDF yields a clear "no extractable text" message.
"""

from __future__ import annotations

import os
from pathlib import Path

from pydantic import BaseModel, ConfigDict, Field

from core.base import Tool, ToolContext, ToolResult
from core.effects import Effect, FileMode, FilesystemEffect
from core.paths import (
    PathResolutionChangedError,
    open_no_follow_read_bytes,
    resolve_real,
)

# Extraction bounds — constants, not config (one big file must not blow
# the context budget; the derived budget from Phase 4 assumes bounded
# tool results).
MAX_EXTRACT_CHARS = 40_000
MAX_DOCUMENT_BYTES = 50 * 1024 * 1024  # refuse before parsing, not after

SCANNED_PDF_MESSAGE = (
    "no extractable text (the PDF looks scanned/image-based; OCR is not "
    "available)"
)


class _PathParams(BaseModel):
    model_config = ConfigDict(extra="forbid")
    path: str = Field(min_length=1)


def _blessed_path(effects: list[Effect]) -> Path:
    for effect in effects:
        if isinstance(effect, FilesystemEffect):
            return Path(effect.path)
    raise ValueError(
        "execute requires the gate-evaluated FilesystemEffect; "
        "it must never run outside enforce_and_run"
    )


def _resolve_path(path: str, context: ToolContext) -> Path:
    p = Path(path)
    if not p.is_absolute() and context.system.sandbox.fs_root is not None:
        p = Path(context.system.sandbox.fs_root) / p
    return resolve_real(p)


def _check_size(target: Path, tool_name: str) -> ToolResult | None:
    size = os.stat(target).st_size
    if size > MAX_DOCUMENT_BYTES:
        return ToolResult(
            ok=False,
            error=(
                f"{tool_name} failed: file is {size} bytes, above the "
                f"{MAX_DOCUMENT_BYTES}-byte limit for document extraction"
            ),
        )
    return None


def _accumulate(pieces, cap: int) -> tuple[str, bool, int]:
    """Join text pieces until ``cap`` chars; stop DURING extraction.

    Returns (text, truncated, pieces_consumed).
    """
    parts: list[str] = []
    total = 0
    consumed = 0
    truncated = False
    for piece in pieces:
        consumed += 1
        if not piece:
            continue
        room = cap - total
        if room <= 0:
            truncated = True
            consumed -= 1
            break
        if len(piece) > room:
            parts.append(piece[:room])
            total += room
            truncated = True
            break
        parts.append(piece)
        total += len(piece)
    return "\n".join(parts), truncated, consumed


class ReadPdfTool(Tool):
    name = "read_pdf"
    description = (
        "Extract the text of a text-based PDF inside the sandbox "
        "(no OCR; long documents are truncated)."
    )
    input_schema = _PathParams
    mutating = False
    destructive = False

    def plan_effects(self, params: _PathParams, context: ToolContext) -> list[Effect]:
        return [FilesystemEffect(path=str(_resolve_path(params.path, context)), mode=FileMode.READ)]

    def execute(
        self, params: _PathParams, context: ToolContext, effects: list[Effect]
    ) -> ToolResult:
        from pypdf import PdfReader  # sealed here; imported at use

        target = _blessed_path(effects)
        try:
            too_big = _check_size(target, self.name)
            if too_big is not None:
                return too_big
            with open_no_follow_read_bytes(target) as f:
                reader = PdfReader(f)
                if reader.is_encrypted:
                    return ToolResult(
                        ok=False,
                        error="read_pdf failed: the PDF is password-protected/"
                              "encrypted and cannot be read",
                    )
                n_pages = len(reader.pages)
                text, truncated, pages_used = _accumulate(
                    (page.extract_text() or "" for page in reader.pages),
                    MAX_EXTRACT_CHARS,
                )
            if not text.strip():
                return ToolResult(ok=True, output=SCANNED_PDF_MESSAGE)
            if truncated:
                text += (
                    f"\n[document truncated: showing first {len(text)} "
                    f"characters from {pages_used} of {n_pages} pages]"
                )
            return ToolResult(ok=True, output=text)
        except (OSError, PathResolutionChangedError) as e:
            return ToolResult(ok=False, error=f"read_pdf failed: {type(e).__name__}: {e}")
        except Exception as e:  # malformed PDFs raise library-specific errors
            return ToolResult(
                ok=False,
                error=f"read_pdf failed: could not parse the PDF "
                      f"({type(e).__name__}: {e})",
            )


def _docx_blocks(document):
    """Body text as plain-text blocks: paragraphs (headings included),
    then tables as tab-separated rows."""
    for paragraph in document.paragraphs:
        yield paragraph.text
    for table in document.tables:
        for row in table.rows:
            yield "\t".join(cell.text for cell in row.cells)


class ReadDocxTool(Tool):
    name = "read_docx"
    description = (
        "Extract the body text of a Word .docx inside the sandbox "
        "(paragraphs, headings, tables as plain text; long documents are "
        "truncated)."
    )
    input_schema = _PathParams
    mutating = False
    destructive = False

    def plan_effects(self, params: _PathParams, context: ToolContext) -> list[Effect]:
        return [FilesystemEffect(path=str(_resolve_path(params.path, context)), mode=FileMode.READ)]

    def execute(
        self, params: _PathParams, context: ToolContext, effects: list[Effect]
    ) -> ToolResult:
        from docx import Document  # sealed here; imported at use

        target = _blessed_path(effects)
        try:
            too_big = _check_size(target, self.name)
            if too_big is not None:
                return too_big
            with open_no_follow_read_bytes(target) as f:
                document = Document(f)
                blocks = list(_docx_blocks(document))
            text, truncated, blocks_used = _accumulate(blocks, MAX_EXTRACT_CHARS)
            if truncated:
                text += (
                    f"\n[document truncated: showing first {len(text)} "
                    f"characters from {blocks_used} of {len(blocks)} blocks]"
                )
            return ToolResult(ok=True, output=text)
        except (OSError, PathResolutionChangedError) as e:
            return ToolResult(ok=False, error=f"read_docx failed: {type(e).__name__}: {e}")
        except Exception as e:  # malformed/corrupt .docx (e.g. BadZipFile)
            return ToolResult(
                ok=False,
                error=f"read_docx failed: could not parse the document "
                      f"({type(e).__name__}: {e})",
            )
